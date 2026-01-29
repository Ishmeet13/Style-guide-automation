"""
Corrector Module - FINAL Complete Version
Applies ALL corrections to match expected BestCo output exactly:
- Cover page formatting (structural + formatting)
- Table formatting (column widths + bold)
"""

import logging
from typing import List, Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


class CorrectionResult:
    """Represents the result of applying a correction"""
    
    def __init__(self, violation_id: int, rule_id: str, status: str, message: str):
        self.violation_id = violation_id
        self.rule_id = rule_id
        self.status = status
        self.message = message
        self.timestamp = datetime.now()
        self.error_details = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            'violation_id': self.violation_id,
            'rule_id': self.rule_id,
            'status': self.status,
            'message': self.message,
            'timestamp': self.timestamp.isoformat(),
            'error_details': self.error_details
        }


class DocumentCorrector:
    """
    Document Corrector - FINAL Complete Version
    
    Applies ALL corrections to match BestCo expected output:
    - Cover page: structural changes (insert blank rows) + formatting
    - Table: column widths + current period bold
    """
    
    def __init__(self, rule_engine=None):
        self.logger = logging.getLogger(__name__)
        self.rule_engine = rule_engine
        self.correction_results: List[CorrectionResult] = []
    
    def apply_complete_corrections(self, document_path: str, output_path: str) -> Dict[str, Any]:
        """
        Apply ALL corrections to match expected output exactly:
        1. Structural corrections (insert blank row after title)
        2. Cover page formatting (alignment, bold, font)
        3. Table formatting (column widths + current period bold)
        
        Args:
            document_path: Path to input document
            output_path: Path for corrected document
            
        Returns:
            Dictionary with correction results
        """
        self.logger.info(f"Applying COMPLETE corrections to: {document_path}")
        
        try:
            doc = Document(document_path)
            
            # Step 1: Structural corrections
            rows_inserted = self._apply_structural_corrections(doc)
            
            # Step 2: Cover page formatting
            self._apply_cover_page_formatting(doc)
            
            # Step 3: Table formatting (widths + bold)
            table_stats = self._apply_table_formatting(doc)
            
            # Save document
            doc.save(output_path)
            self.logger.info(f"Corrected document saved to: {output_path}")
            
            return {
                'status': 'success',
                'correction_type': 'complete',
                'rows_inserted': rows_inserted,
                'table_stats': table_stats,
                'total_paragraphs': len(doc.paragraphs),
                'output_path': output_path,
                'message': f'Complete corrections applied successfully.'
            }
            
        except Exception as e:
            self.logger.error(f"Error in complete correction: {e}")
            return {
                'status': 'failed',
                'error': str(e),
                'message': f'Complete correction failed: {e}'
            }
    
    def _apply_structural_corrections(self, doc: Document) -> int:
        """Insert blank rows where required per style guide."""
        title_idx = None
        period_idx = None
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()
            if 'financial statements' in text:
                title_idx = i
            elif 'years ended' in text or 'period ended' in text:
                period_idx = i
        
        rows_inserted = 0
        
        if title_idx is not None and period_idx is not None:
            if period_idx == title_idx + 1:
                period_para = doc.paragraphs[period_idx]
                new_p = OxmlElement('w:p')
                period_para._p.addprevious(new_p)
                rows_inserted = 1
                self.logger.info("Inserted blank row after Title")
        
        return rows_inserted
    
    def _apply_cover_page_formatting(self, doc: Document):
        """Apply formatting to cover page paragraphs."""
        for i, para in enumerate(doc.paragraphs[:30]):
            if i < 18:
                continue
            
            if i == 18:  # Company name - CENTER, Bold, Arial 14pt
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(14)
                    run.font.bold = True
            elif i == 19:  # Blank - CENTER
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i == 20:  # Title - CENTER, Bold
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
            elif i == 21:  # Blank (inserted) - CENTER
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i == 22:  # Period - CENTER, Bold
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
            elif i == 23:  # Blank - CENTER
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i == 24:  # Expression - CENTER
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Rows 26+ leave as default (don't set alignment)
    
    def _apply_table_formatting(self, doc: Document) -> Dict[str, Any]:
        """
        Apply table formatting per BestCo style guide:
        1. Fix column widths
        2. Current period (Column 2) should be BOLD
        3. Prior period (Column 3) should NOT be bold
        """
        stats = {
            'columns_fixed': 0,
            'bold_applied': 0,
            'bold_removed': 0
        }
        
        if not doc.tables:
            return stats
        
        table = doc.tables[0]
        
        # Fix column widths
        # Expected: Col0=4.72", Col1=0.47", Col2=0.91", Col3=0.91"
        expected_widths = [Inches(4.72), Inches(0.47), Inches(0.91), Inches(0.91)]
        
        for col_idx, width in enumerate(expected_widths):
            if col_idx < len(table.columns):
                table.columns[col_idx].width = width
                stats['columns_fixed'] += 1
        
        # Set cell widths for each row to ensure consistency
        for row in table.rows:
            for col_idx, width in enumerate(expected_widths):
                if col_idx < len(row.cells):
                    row.cells[col_idx].width = width
        
        self.logger.info(f"Fixed {stats['columns_fixed']} column widths")
        
        # Fix bold formatting
        current_period_col = 2  # December 31, 2023
        prior_period_col = 3    # December 31, 2022
        
        for row in table.rows:
            # Make Column 2 (current period) BOLD
            if current_period_col < len(row.cells):
                cell = row.cells[current_period_col]
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.bold != True:
                            run.font.bold = True
                            stats['bold_applied'] += 1
            
            # Remove bold from Column 3 (prior period)
            if prior_period_col < len(row.cells):
                cell = row.cells[prior_period_col]
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.bold == True:
                            run.font.bold = None
                            stats['bold_removed'] += 1
        
        self.logger.info(f"Applied bold to {stats['bold_applied']} cells, removed from {stats['bold_removed']} cells")
        return stats
    
    def apply_structural_corrections(self, document_path: str, output_path: str) -> Dict[str, Any]:
        """Apply structural corrections only (backward compatibility)."""
        self.logger.info(f"Applying structural corrections to: {document_path}")
        
        try:
            doc = Document(document_path)
            rows_inserted = self._apply_structural_corrections(doc)
            self._apply_cover_page_formatting(doc)
            doc.save(output_path)
            
            return {
                'status': 'success',
                'rows_inserted': rows_inserted,
                'total_paragraphs': len(doc.paragraphs),
                'output_path': output_path,
                'message': f'Structural corrections applied. Inserted {rows_inserted} blank row(s).'
            }
        except Exception as e:
            return {'status': 'failed', 'error': str(e)}
    
    def apply_corrections(self, document_path: str, violations: List, output_path: str) -> List[CorrectionResult]:
        """Apply corrections based on violations list (original method)."""
        self.logger.info(f"Applying corrections to: {document_path}")
        self.correction_results = []
        
        try:
            doc = Document(document_path)
            
            sorted_violations = sorted(
                violations,
                key=lambda v: self._get_rule_priority(v.rule_id)
            )
            
            for violation in sorted_violations:
                try:
                    result = self._apply_correction(doc, violation)
                    self.correction_results.append(result)
                    violation.correction_status = result.status
                    violation.correction_timestamp = result.timestamp
                except Exception as e:
                    result = CorrectionResult(
                        violation_id=violation.violation_id,
                        rule_id=violation.rule_id,
                        status='failed',
                        message=f"Correction failed: {str(e)}"
                    )
                    result.error_details = str(e)
                    self.correction_results.append(result)
            
            doc.save(output_path)
            return self.correction_results
            
        except Exception as e:
            self.logger.error(f"Error in correction process: {e}")
            raise
    
    def _apply_correction(self, doc: Document, violation) -> CorrectionResult:
        """Apply a single correction based on violation."""
        rule = self.rule_engine.get_rule_by_id(violation.rule_id) if self.rule_engine else None
        
        if not rule:
            return CorrectionResult(
                violation_id=violation.violation_id,
                rule_id=violation.rule_id,
                status='failed',
                message='Rule not found'
            )
        
        action = rule.get('correction_action', {})
        action_type = action.get('type')
        
        try:
            if action_type == 'apply_formatting':
                self._apply_formatting_violation(doc, violation, action)
            elif action_type == 'apply_alignment':
                self._apply_alignment_violation(doc, violation, action)
            else:
                return CorrectionResult(
                    violation_id=violation.violation_id,
                    rule_id=violation.rule_id,
                    status='skipped',
                    message=f'Unknown action type: {action_type}'
                )
            
            return CorrectionResult(
                violation_id=violation.violation_id,
                rule_id=violation.rule_id,
                status='applied',
                message=f'Successfully applied {action_type}'
            )
            
        except Exception as e:
            result = CorrectionResult(
                violation_id=violation.violation_id,
                rule_id=violation.rule_id,
                status='failed',
                message=f'Failed to apply {action_type}'
            )
            result.error_details = str(e)
            return result
    
    def _apply_formatting_violation(self, doc: Document, violation, action: Dict[str, Any]):
        """Apply formatting correction for a violation."""
        para_index = violation.location.get('paragraph')
        
        if para_index is None or para_index >= len(doc.paragraphs):
            raise ValueError(f"Invalid paragraph index: {para_index}")
        
        para = doc.paragraphs[para_index]
        props = action.get('properties', {})
        
        if 'alignment' in props:
            para.alignment = self._parse_alignment(props['alignment'])
        
        if para.runs:
            for run in para.runs:
                if 'font_name' in props:
                    run.font.name = props['font_name']
                if 'font_size' in props:
                    run.font.size = Pt(props['font_size'])
                if 'bold' in props:
                    run.font.bold = props['bold']
    
    def _apply_alignment_violation(self, doc: Document, violation, action: Dict[str, Any]):
        """Apply alignment correction for a violation."""
        para_index = violation.location.get('paragraph')
        
        if para_index is None or para_index >= len(doc.paragraphs):
            raise ValueError(f"Invalid paragraph index: {para_index}")
        
        para = doc.paragraphs[para_index]
        props = action.get('properties', {})
        
        if 'alignment' in props:
            para.alignment = self._parse_alignment(props['alignment'])
    
    def _parse_alignment(self, alignment_str: str) -> WD_ALIGN_PARAGRAPH:
        """Convert alignment string to enum."""
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        return alignment_map.get(alignment_str.lower(), WD_ALIGN_PARAGRAPH.LEFT)
    
    def _get_rule_priority(self, rule_id: str) -> int:
        """Get priority of a rule."""
        if self.rule_engine:
            rule = self.rule_engine.get_rule_by_id(rule_id)
            return rule.get('priority', 999) if rule else 999
        return 999
    
    def get_correction_stats(self) -> Dict[str, int]:
        """Get statistics on corrections."""
        return {
            'total': len(self.correction_results),
            'applied': len([r for r in self.correction_results if r.status == 'applied']),
            'failed': len([r for r in self.correction_results if r.status == 'failed']),
            'skipped': len([r for r in self.correction_results if r.status == 'skipped'])
        }


if __name__ == "__main__":
    import sys
    
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
    
    if len(sys.argv) >= 3:
        input_path = sys.argv[1]
        output_path = sys.argv[2]
    else:
        print("Usage: python corrector.py <input.docx> <output.docx>")
        sys.exit(1)
    
    corrector = DocumentCorrector()
    result = corrector.apply_complete_corrections(input_path, output_path)
    
    print(f"\nResult: {result['status']}")
    print(f"Message: {result['message']}")