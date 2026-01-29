"""
Validator Module
Validates documents against style guide rules and detects violations

"""

import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Violation:
    """Represents a single rule violation"""
    
    def __init__(self, rule_id: str, rule_name: str, severity: str,
                 location: Dict[str, Any], expected: Dict[str, Any],
                 actual: Dict[str, Any], message: str):
        self.violation_id = None  # Will be set by validator
        self.rule_id = rule_id
        self.rule_name = rule_name
        self.severity = severity
        self.location = location
        self.expected = expected
        self.actual = actual
        self.message = message
        self.correction_status = "pending"
        self.correction_timestamp = None
        
    def to_dict(self) -> Dict[str, Any]:
        """Convert violation to dictionary"""
        return {
            'violation_id': self.violation_id,
            'rule_id': self.rule_id,
            'rule_name': self.rule_name,
            'severity': self.severity,
            'location': self.location,
            'expected': self.expected,
            'actual': self.actual,
            'message': self.message,
            'correction_status': self.correction_status,
            'correction_timestamp': self.correction_timestamp
        }


class DocumentValidator:
    """
    Document Validator
    
    Validates .docx documents against style guide rules
    Detects formatting violations and generates detailed reports
    """
    
    def __init__(self, rule_engine):
        """
        Initialize Validator
        
        Args:
            rule_engine: RuleEngine instance with loaded rules
        """
        self.logger = logging.getLogger(__name__)
        self.rule_engine = rule_engine
        self.violations: List[Violation] = []
        self.violation_counter = 0
        
    def validate_document(self, document_path: str) -> List[Violation]:
        """
        Validate a document against all enabled rules
        
        Args:
            document_path: Path to .docx file
            
        Returns:
            List of violations found
        """
        self.logger.info(f"Starting validation of: {document_path}")
        self.violations = []
        self.violation_counter = 0
        
        try:
            # Load document
            doc = Document(document_path)
            
            # Get enabled rules
            rules = self.rule_engine.get_enabled_rules()
            self.logger.info(f"Validating against {len(rules)} enabled rules")
            
            # Validate each rule
            for rule in rules:
                try:
                    self._validate_rule(doc, rule)
                except Exception as e:
                    self.logger.error(f"Error validating rule {rule.get('rule_id')}: {e}")
            
            self.logger.info(f"Validation complete. Found {len(self.violations)} violations")
            return self.violations
            
        except Exception as e:
            self.logger.error(f"Error validating document: {e}")
            raise
    
    def _validate_rule(self, doc: Document, rule: Dict[str, Any]):
        """Validate a single rule"""
        category = rule.get('category')
        
        if category == 'cover_page':
            self._validate_cover_page_rule(doc, rule)
        elif category in ['table_structure', 'table_formatting']:
            self._validate_table_rule(doc, rule)
        else:
            self.logger.debug(f"Skipping unknown category: {category}")
    
    def _validate_cover_page_rule(self, doc: Document, rule: Dict[str, Any]):
        """Validate cover page rules"""
        rule_id = rule.get('rule_id')
        location = rule.get('location', {})
        validation = rule.get('validation', {})
        
        # Get row/paragraph to check
        row_from_top = location.get('row_from_top')
        
        if row_from_top is None:
            return
        
        # Handle single row number
        if isinstance(row_from_top, int):
            para_index = row_from_top - 1
            self._validate_paragraph(doc, para_index, rule)
        elif isinstance(row_from_top, str):
            if row_from_top.isdigit():
                para_index = int(row_from_top) - 1
                self._validate_paragraph(doc, para_index, rule)
            elif '-' in row_from_top:
                # Handle range (e.g., "1-18")
                start, end = map(int, row_from_top.split('-'))
                for para_index in range(start - 1, min(end, len(doc.paragraphs))):
                    self._validate_paragraph(doc, para_index, rule)
            elif "all" in row_from_top.lower():
                # Handle "all" selector
                for i in range(min(30, len(doc.paragraphs))):
                    self._validate_paragraph(doc, i, rule)
    
    def _validate_paragraph(self, doc: Document, para_index: int, rule: Dict[str, Any]):
        """Validate a specific paragraph"""
        if para_index >= len(doc.paragraphs):
            return
        
        para = doc.paragraphs[para_index]
        validation = rule.get('validation', {})
        violations_found = False
        expected = {}
        actual = {}
        
        # Check alignment - CRITICAL: None or 0 means LEFT
        if 'alignment' in validation:
            expected_align_str = validation['alignment'].lower()
            expected_align_val = self._alignment_str_to_int(expected_align_str)
            actual_align_val = self._get_alignment_as_int(para.alignment)
            
            expected['alignment'] = expected_align_str
            actual['alignment'] = self._alignment_int_to_str(actual_align_val)
            
            if expected_align_val != actual_align_val:
                violations_found = True
                self.logger.debug(f"Row {para_index + 1}: Alignment mismatch - expected {expected_align_str}, got {actual['alignment']}")
        
        # Check bold - works for both empty and non-empty paragraphs
        if 'bold' in validation:
            expected_bold = validation['bold']
            actual_bold = self._get_paragraph_bold(para)
            
            expected['bold'] = expected_bold
            actual['bold'] = actual_bold
            
            # Normalize: None is treated as False for comparison
            actual_bold_normalized = actual_bold if actual_bold is not None else False
            
            if expected_bold != actual_bold_normalized:
                violations_found = True
                self.logger.debug(f"Row {para_index + 1}: Bold mismatch - expected {expected_bold}, got {actual_bold}")
        
        # Check font properties (only if runs exist)
        if para.runs:
            run = para.runs[0]
            
            # Font name
            if 'font_name' in validation:
                expected['font_name'] = validation['font_name']
                actual_font = run.font.name
                actual['font_name'] = actual_font or "default"
                
                # Only flag if explicitly different (not None)
                if actual_font is not None and actual_font != validation['font_name']:
                    violations_found = True
            
            # Font size
            if 'font_size' in validation:
                expected['font_size'] = validation['font_size']
                actual_size = run.font.size.pt if run.font.size else None
                actual['font_size'] = actual_size
                
                if actual_size is not None and actual_size != validation['font_size']:
                    violations_found = True
        else:
            # Empty paragraph - check if we need specific formatting
            if 'font_size' in validation:
                expected['font_size'] = validation['font_size']
                actual['font_size'] = None
                # Don't flag as violation for empty paragraphs with no runs
            
            if 'font_name' in validation:
                expected['font_name'] = validation['font_name']
                actual['font_name'] = None
                # Don't flag as violation for empty paragraphs with no runs
        
        # Add violation if any check failed
        if violations_found:
            message = self._generate_message(rule, para_index, expected, actual)
            self._add_violation(
                rule=rule,
                location={'page': 1, 'paragraph': para_index, 'row': para_index + 1},
                expected=expected,
                actual=actual,
                message=message
            )
    
    def _get_alignment_as_int(self, alignment) -> int:
        """Convert alignment to integer (0=LEFT, 1=CENTER, 2=RIGHT, 3=JUSTIFY)"""
        if alignment is None:
            return 0  # None means LEFT
        if isinstance(alignment, int):
            return alignment
        # WD_ALIGN_PARAGRAPH enum values
        try:
            return int(alignment)
        except:
            return 0
    
    def _alignment_str_to_int(self, align_str: str) -> int:
        """Convert alignment string to integer"""
        mapping = {
            'left': 0,
            'center': 1,
            'right': 2,
            'justify': 3
        }
        return mapping.get(align_str.lower(), 0)
    
    def _alignment_int_to_str(self, align_int: int) -> str:
        """Convert alignment integer to string"""
        mapping = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
        return mapping.get(align_int, 'left')
    
    def _get_paragraph_bold(self, para) -> Optional[bool]:
        """Get bold status of paragraph, handling empty paragraphs"""
        if para.runs:
            return para.runs[0].font.bold
        return None
    
    def _validate_table_rule(self, doc: Document, rule: Dict[str, Any]):
        """Validate table rules"""
        if not doc.tables:
            return
        
        rule_id = rule.get('rule_id')
        validation = rule.get('validation', {})
        
        for table_idx, table in enumerate(doc.tables):
            
            # Validate row heights
            if rule_id == 'TABLE_ROW_HEIGHT':
                expected_height = validation.get('row_height', 0.37)
                
                for row_idx, row in enumerate(table.rows):
                    if row.height:
                        actual_height = row.height.cm
                        
                        # Allow 0.05cm tolerance
                        if abs(actual_height - expected_height) > 0.05:
                            self._add_violation(
                                rule=rule,
                                location={'table': table_idx, 'row': row_idx},
                                expected={'row_height_cm': expected_height},
                                actual={'row_height_cm': round(actual_height, 2)},
                                message=f"Table {table_idx} Row {row_idx}: height {actual_height:.2f}cm (expected {expected_height}cm)"
                            )
            
            # Validate column widths
            elif rule_id == 'TABLE_VALUE_COLUMN_WIDTH':
                expected_width = validation.get('column_width', 2.3)
                
                # Check last 2-3 columns (typically value columns)
                num_cols = len(table.columns)
                for col_idx in range(max(0, num_cols - 3), num_cols):
                    col = table.columns[col_idx]
                    if col.width:
                        actual_width = col.width.cm
                        
                        if abs(actual_width - expected_width) > 0.1:
                            self._add_violation(
                                rule=rule,
                                location={'table': table_idx, 'column': col_idx},
                                expected={'column_width_cm': expected_width},
                                actual={'column_width_cm': round(actual_width, 2)},
                                message=f"Table {table_idx} Column {col_idx}: width {actual_width:.2f}cm (expected {expected_width}cm)"
                            )
            
            # Validate current period bold
            elif rule_id == 'BALANCE_SHEET_CURRENT_PERIOD_BOLD':
                # Check last column (current period)
                for row_idx, row in enumerate(table.rows[1:], start=1):  # Skip header
                    if len(row.cells) > 0:
                        last_cell = row.cells[-1]
                        
                        for para in last_cell.paragraphs:
                            for run in para.runs:
                                # Check if contains numbers or $
                                if any(c.isdigit() or c == '$' for c in run.text):
                                    if not run.font.bold:
                                        self._add_violation(
                                            rule=rule,
                                            location={'table': table_idx, 'row': row_idx, 'column': len(row.cells) - 1},
                                            expected={'bold': True},
                                            actual={'bold': False},
                                            message=f"Table {table_idx} Row {row_idx}: current period values should be bold"
                                        )
                                        break
    
    def _add_violation(self, rule: Dict[str, Any], location: Dict[str, Any],
                      expected: Dict[str, Any], actual: Dict[str, Any], message: str):
        """Add a violation to the list"""
        self.violation_counter += 1
        
        violation = Violation(
            rule_id=rule.get('rule_id'),
            rule_name=rule.get('description', rule.get('rule_id')),
            severity=rule.get('severity', 'medium'),
            location=location,
            expected=expected,
            actual=actual,
            message=message
        )
        violation.violation_id = self.violation_counter
        
        self.violations.append(violation)
        self.logger.debug(f"Violation #{self.violation_counter}: {rule.get('rule_id')}")
    
    def _generate_message(self, rule: Dict[str, Any], para_index: int,
                         expected: Dict, actual: Dict) -> str:
        """Generate human-readable violation message"""
        differences = []
        for key in expected:
            if key in actual and expected[key] != actual[key]:
                differences.append(f"{key}: expected '{expected[key]}', got '{actual[key]}'")
        
        return f"Row {para_index + 1}: {'; '.join(differences)}"
    
    def _parse_alignment(self, alignment_str: str) -> Optional[WD_ALIGN_PARAGRAPH]:
        """Convert alignment string to enum"""
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        return alignment_map.get(alignment_str.lower())
    
    def _alignment_to_string(self, alignment) -> str:
        """Convert alignment enum to string"""
        if alignment is None or alignment == 0:
            return 'left'
        if alignment == WD_ALIGN_PARAGRAPH.CENTER or alignment == 1:
            return 'center'
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT or alignment == 2:
            return 'right'
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY or alignment == 3:
            return 'justify'
        else:
            return 'left'
    
    def get_violations_by_severity(self, severity: str) -> List[Violation]:
        """Get violations filtered by severity"""
        return [v for v in self.violations if v.severity == severity]
    
    def get_violations_by_category(self, category: str) -> List[Violation]:
        """Get violations filtered by category"""
        category_rule_ids = [
            r.get('rule_id') 
            for r in self.rule_engine.get_rules_by_category(category)
        ]
        return [v for v in self.violations if v.rule_id in category_rule_ids]
    
    def get_summary(self) -> Dict[str, Any]:
        """Get validation summary"""
        return {
            'total_violations': len(self.violations),
            'high_severity': len(self.get_violations_by_severity('high')),
            'medium_severity': len(self.get_violations_by_severity('medium')),
            'low_severity': len(self.get_violations_by_severity('low')),
            'by_category': {
                category: len(self.get_violations_by_category(category))
                for category in self.rule_engine.rules_by_category.keys()
            }
        }


# Example usage
if __name__ == "__main__":
    logging.basicConfig(level=logging.DEBUG)
    
    from rule_engine import RuleEngine
    
    # Initialize components
    engine = RuleEngine("bestco-rules.json")
    validator = DocumentValidator(engine)
    
    # Validate document
    violations = validator.validate_document("bestco-sample-input.docx")
    
    # Print results
    print(f"\nFound {len(violations)} violations\n")
    
    for v in violations[:10]:  # Show first 10
        print(f"[{v.severity.upper()}] {v.rule_name}")
        print(f"  Location: Row {v.location.get('row', 'N/A')}")
        print(f"  Expected: {v.expected}")
        print(f"  Actual: {v.actual}")
        print(f"  Message: {v.message}")
        print()
    
    # Print summary
    summary = validator.get_summary()
    print(f"Summary: {summary}")